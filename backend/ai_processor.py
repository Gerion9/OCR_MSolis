"""
Procesador de IA para generar Declaration Letters
Integración con múltiples proveedores: Google Gemini, Groq, etc.
"""

import os
import time
import logging
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Generator
import google.generativeai as genai
from pathlib import Path

# Configurar logging
logger = logging.getLogger(__name__)

GROQ_MAX_TOKENS_DECLARATION = 16000
GROQ_MAX_TOKENS_COVER = 8000

# Importar Groq
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False
    logger.warning("Groq SDK no disponible. Instale con: pip install groq")

# CONSTANTES DE CONFIGURACIÓN 
# Configuración base de generación para Gemini
DEFAULT_GENERATION_CONFIG = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8000,
}

# Configuración optimizada para Cover Letter (documentos más largos) - Gemini
COVER_LETTER_GENERATION_CONFIG = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 12000,
    "candidate_count": 1,
}


# Configuración de seguridad (necesario para documentos de tráfico humano)
SAFETY_SETTINGS = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_NONE"
    },
]

class BaseAIProcessor:
    """
    Clase base para procesadores de IA
    Define la interfaz común para todos los proveedores
    """
    
    def __init__(self, api_key: str, model_name: str, request_timeout: int = 300):
        self.api_key = api_key
        self.model_name = model_name
        self.request_timeout = request_timeout
        self.system_prompt = ""
        self.declaration_guide = ""
        self.cover_letter_system_prompt = ""
        self.cover_letter_structure = ""
    
    def generate_declaration_letter(self, questionnaire_text: str) -> Optional[str]:
        raise NotImplementedError
    
    def generate_declaration_letter_stream(self, questionnaire_text: str) -> Generator[str, None, None]:
        raise NotImplementedError
    
    def generate_cover_letter(self, declaration_letter_content: str) -> Optional[str]:
        raise NotImplementedError
    
    def generate_cover_letter_stream(self, declaration_letter_content: str) -> Generator[str, None, None]:
        raise NotImplementedError
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        raise NotImplementedError


class GeminiAIProcessor(BaseAIProcessor):
    """
    Procesador de IA para generar declaration letters usando Google Gemini
    """
    
    def __init__(self, api_key: str, model_name: str, request_timeout: int = 300):
        """
        Inicializa el procesador de IA con Gemini
        
        Args:
            api_key: API key de Google Gemini
            model_name: Nombre del modelo a usar
            request_timeout: Timeout en segundos para las solicitudes (default: 300s = 5 minutos)
        """
        super().__init__(api_key, model_name, request_timeout)
        
        # Configurar Gemini
        genai.configure(api_key=api_key)
        
        # Usar configuraciones desde constantes
        self.generation_config = DEFAULT_GENERATION_CONFIG.copy()
        self.cover_letter_generation_config = COVER_LETTER_GENERATION_CONFIG.copy()
        self.safety_settings = SAFETY_SETTINGS.copy()
        
        # Inicializar el modelo
        self.model = genai.GenerativeModel(
            model_name=model_name,
            generation_config=self.generation_config,
            safety_settings=self.safety_settings
        )
        
        # Modelo optimizado para Cover Letter
        self.cover_letter_model = genai.GenerativeModel(
            model_name=model_name,
            generation_config=self.cover_letter_generation_config,
            safety_settings=self.safety_settings
        )
        
        logging.debug(f"Procesador de IA inicializado con modelo: {model_name}")
        logging.debug(f"Timeout configurado: {request_timeout} segundos")
    
    def load_xml_files(self, system_prompt_path: str, declaration_path: str) -> bool:
        """
        Carga los archivos XML con las instrucciones y estructura
        
        Args:
            system_prompt_path: Ruta al archivo SystemPrompt.xml
            declaration_path: Ruta al archivo Declaration.xml
        
        Returns:
            bool: True si se cargaron correctamente
        """
        try:
            # Leer SystemPrompt.xml
            with open(system_prompt_path, 'r', encoding='utf-8') as f:
                self.system_prompt = f.read()
            
            # Leer Declaration.xml
            with open(declaration_path, 'r', encoding='utf-8') as f:
                self.declaration_guide = f.read()
            
            logging.debug("Archivos XML de Declaration Letter cargados correctamente")
            return True
        
        except Exception as e:
            logging.debug(f"Error al cargar archivos XML de Declaration Letter: {e}")
            return False
    
    def load_cover_letter_xml_files(self, system_prompt_path: str, structure_path: str) -> bool:
        """
        Carga los archivos XML para Cover Letter
        
        Args:
            system_prompt_path: Ruta al archivo SystemPrompt.xml de Cover Letter
            structure_path: Ruta al archivo CoverLetterStructure.xml
        
        Returns:
            bool: True si se cargaron correctamente
        """
        try:
            # Leer SystemPrompt.xml de Cover Letter
            with open(system_prompt_path, 'r', encoding='utf-8') as f:
                self.cover_letter_system_prompt = f.read()
            
            # Leer CoverLetterStructure.xml
            with open(structure_path, 'r', encoding='utf-8') as f:
                self.cover_letter_structure = f.read()
            
            logging.debug("Archivos XML de Cover Letter cargados correctamente")
            return True
        
        except Exception as e:
            logger.error(f"Error al cargar archivos XML de Cover Letter: {e}")
            return False
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extrae texto de un archivo
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            str: Texto extraído o None si hay error
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Para archivos de texto plano
            if file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            # Para archivos DOCX
            elif file_extension == '.docx':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = []
                    for paragraph in doc.paragraphs:
                        text.append(paragraph.text)
                    return '\n'.join(text)
                except ImportError:
                    # Si python-docx no está disponible, intentar lectura básica
                    logger.warning("python-docx no disponible, usando lectura básica")
                    return self._extract_text_from_docx_basic(file_path)
            
            # Para archivos PDF
            elif file_extension == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = []
                        for page in pdf_reader.pages:
                            text.append(page.extract_text())
                        return '\n'.join(text)
                except ImportError:
                    logger.warning("PyPDF2 no disponible")
                    return None
            
            else:
                logger.error(f"Tipo de archivo no soportado: {file_extension}")
                return None
        
        except Exception as e:
            logger.error(f"Error al extraer texto: {e}")
            return None
    
    def _extract_text_from_docx_basic(self, file_path: str) -> Optional[str]:
        """
        Extrae texto de un DOCX usando solo librerías estándar
        
        Args:
            file_path: Ruta al archivo DOCX
        
        Returns:
            str: Texto extraído
        """
        try:
            import zipfile
            
            with zipfile.ZipFile(file_path, 'r') as docx:
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Extraer todo el texto
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                paragraphs = tree.findall('.//w:p', namespaces)
                
                text = []
                for paragraph in paragraphs:
                    texts = paragraph.findall('.//w:t', namespaces)
                    paragraph_text = ''.join([t.text for t in texts if t.text])
                    if paragraph_text:
                        text.append(paragraph_text)
                
                return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error en extracción básica de DOCX: {e}")
            return None
    
    # MÉTODOS PRIVADOS DE GENERACIÓN 
    def _generate_content(
        self, 
        prompt: str, 
        model: genai.GenerativeModel, 
        doc_type: str = "documento",
        use_stream: bool = False
    ):
        """
        Método privado unificado para generar contenido (normal o streaming)
        
        Args:
            prompt: Prompt completo
            model: Modelo de Gemini a usar
            doc_type: Tipo de documento para logging
            use_stream: Si usar streaming o no
            
        Returns/Yields:
            str o Generator de strings
        """
        try:
            logging.debug(f"Generando {doc_type} con IA...")
            logging.debug(f"Usando timeout de {self.request_timeout} segundos...")
            
            start_time = time.time()
            
            if use_stream:
                # Modo streaming
                response = model.generate_content(prompt, stream=True)
                
                for chunk in response:
                    if chunk.text:
                        yield chunk.text
                
                elapsed_time = time.time() - start_time
                logging.debug(f"Generación con streaming de {doc_type} completada en {elapsed_time:.2f} segundos")
        
        except Exception as e:
            error_msg = str(e)
            if "timeout" in error_msg.lower() or "timed out" in error_msg.lower() or "ReadTimeout" in str(type(e).__name__):
                logger.error(f"La generación de {doc_type} excedió el tiempo límite de {self.request_timeout}s")
                logging.debug("Sugerencia: El documento es muy largo o el servidor está ocupado. Intente nuevamente.")
            else:
                logger.error(f"Error al generar {doc_type}: {e}")
            raise Exception(f"Error al generar {doc_type}: {error_msg}")
    
    # MÉTODOS PÚBLICOS DE GENERACIÓN
    def generate_declaration_letter(self, questionnaire_text: str) -> Optional[str]:
        """
        Genera una declaration letter basada en el cuestionario
        
        Args:
            questionnaire_text: Texto del cuestionario del afectado
        
        Returns:
            str: Declaration letter en formato Markdown o None si hay error
        """
        full_prompt = self._build_prompt(questionnaire_text)
        return self._generate_content(full_prompt, self.model, "Declaration Letter", use_stream=False)
    
    def _build_prompt(self, questionnaire_text: str) -> str:
        """
        Construye el prompt completo para la IA
        
        Args:
            questionnaire_text: Texto del cuestionario
        
        Returns:
            str: Prompt completo
        """
        prompt = f"""
{self.system_prompt}

{self.declaration_guide}

---

CUESTIONARIO DEL AFECTADO:
{questionnaire_text}

---
"""
        return prompt
    
    def generate_cover_letter(self, declaration_letter_content: str) -> Optional[str]:
        """
        Genera un Cover Letter basado en el Declaration Letter
        
        Args:
            declaration_letter_content: Contenido completo del Declaration Letter
        
        Returns:
            str: Cover Letter en formato Markdown o None si hay error
        """
        # Validar que se hayan cargado los archivos XML de Cover Letter
        if not self.cover_letter_system_prompt or not self.cover_letter_structure:
            logger.error("Archivos XML de Cover Letter no cargados")
            return None
        
        full_prompt = self._build_cover_letter_prompt(declaration_letter_content)
        return self._generate_content(full_prompt, self.cover_letter_model, "Cover Letter", use_stream=False)
    
    def _build_cover_letter_prompt(self, declaration_letter_content: str) -> str:
        """
        Construye el prompt completo para generar el Cover Letter
        
        Args:
            declaration_letter_content: Contenido del Declaration Letter
        
        Returns:
            str: Prompt completo
        """
        prompt = f"""
{self.cover_letter_system_prompt}

{self.cover_letter_structure}

---

DECLARATION LETTER DEL SOBREVIVIENTE:
{declaration_letter_content}

---
"""
        return prompt
    
    def generate_declaration_letter_stream(self, questionnaire_text: str) -> Generator[str, None, None]:
        """
        Genera una declaration letter basada en el cuestionario usando streaming
        
        Args:
            questionnaire_text: Texto del cuestionario del afectado
            
        Yields:
            str: Chunks de texto generados en tiempo real
        """
        full_prompt = self._build_prompt(questionnaire_text)
        yield from self._generate_content(full_prompt, self.model, "Declaration Letter", use_stream=True)
    
    def generate_cover_letter_stream(self, declaration_letter_content: str) -> Generator[str, None, None]:
        """
        Genera un Cover Letter basado en el Declaration Letter usando streaming
        
        Args:
            declaration_letter_content: Contenido completo del Declaration Letter
        
        Yields:
            str: Chunks de texto generados en tiempo real
        """
        # Validar que se hayan cargado los archivos XML de Cover Letter
        if not self.cover_letter_system_prompt or not self.cover_letter_structure:
            logger.error("Archivos XML de Cover Letter no cargados")
            raise Exception("Cover Letter XML files not loaded")
        
        full_prompt = self._build_cover_letter_prompt(declaration_letter_content)
        yield from self._generate_content(full_prompt, self.cover_letter_model, "Cover Letter", use_stream=True)
    
    def validate_api_key(self) -> bool:
        """
        Valida que la API key funcione correctamente
        
        Returns:
            bool: True si la API key es válida
        """
        try:
            # Intenta generar un texto simple para validar
            test_model = genai.GenerativeModel(model_name=self.model_name)
            response = test_model.generate_content("Hello")
            logging.debug("API key validada exitosamente")
            return response is not None
        except Exception as e:
            logger.error(f"Error al validar API key: {e}")
            return False

# Mantener compatibilidad con código antiguo
AIProcessor = GeminiAIProcessor


# ==================== GROQ AI PROCESSOR ====================
class GroqAIProcessor(BaseAIProcessor):
    """
    Procesador de IA para generar declaration letters usando Groq
    """
    
    def __init__(self, api_key: str, model_name: str, request_timeout: int = 300):
        """
        Inicializa el procesador de IA con Groq
        
        Args:
            api_key: API key de Groq
            model_name: Nombre del modelo a usar (ej: llama-3.3-70b-versatile)
            request_timeout: Timeout en segundos para las solicitudes
        """
        super().__init__(api_key, model_name, request_timeout)
        
        if not GROQ_AVAILABLE:
            raise ImportError("Groq SDK no está instalado. Instale con: pip install groq")
        
        # Inicializar cliente de Groq
        self.client = Groq(api_key=api_key)
        
        logger.debug(f"Groq AI Processor inicializado con modelo: {model_name}")
        logger.debug(f"Timeout configurado: {request_timeout} segundos")
    
    def load_xml_files(self, system_prompt_path: str, declaration_path: str) -> bool:
        """Carga archivos XML de Declaration Letter"""
        try:
            with open(system_prompt_path, 'r', encoding='utf-8') as f:
                self.system_prompt = f.read()
            
            with open(declaration_path, 'r', encoding='utf-8') as f:
                self.declaration_guide = f.read()
            
            logger.debug("Archivos XML de Declaration Letter cargados correctamente")
            return True
        except Exception as e:
            logger.error(f"Error al cargar archivos XML de Declaration Letter: {e}")
            return False
    
    def load_cover_letter_xml_files(self, system_prompt_path: str, structure_path: str) -> bool:
        """Carga archivos XML de Cover Letter"""
        try:
            with open(system_prompt_path, 'r', encoding='utf-8') as f:
                self.cover_letter_system_prompt = f.read()
            
            with open(structure_path, 'r', encoding='utf-8') as f:
                self.cover_letter_structure = f.read()
            
            logger.debug("Archivos XML de Cover Letter cargados correctamente")
            return True
        except Exception as e:
            logger.error(f"Error al cargar archivos XML de Cover Letter: {e}")
            return False
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extrae texto de un archivo
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            str: Texto extraído o None si hay error
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Para archivos de texto plano
            if file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            # Para archivos DOCX
            elif file_extension == '.docx':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = []
                    for paragraph in doc.paragraphs:
                        text.append(paragraph.text)
                    return '\n'.join(text)
                except ImportError:
                    logger.warning("python-docx no disponible, usando lectura básica")
                    return self._extract_text_from_docx_basic(file_path)
            
            # Para archivos PDF
            elif file_extension == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = []
                        for page in pdf_reader.pages:
                            text.append(page.extract_text())
                        return '\n'.join(text)
                except ImportError:
                    logger.warning("PyPDF2 no disponible")
                    return None
            
            else:
                logger.error(f"Tipo de archivo no soportado: {file_extension}")
                return None
        
        except Exception as e:
            logger.error(f"Error al extraer texto: {e}")
            return None
    
    def _extract_text_from_docx_basic(self, file_path: str) -> Optional[str]:
        """
        Extrae texto de un DOCX usando solo librerías estándar
        """
        try:
            import zipfile
            
            with zipfile.ZipFile(file_path, 'r') as docx:
                xml_content = docx.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                paragraphs = tree.findall('.//w:p', namespaces)
                
                text = []
                for paragraph in paragraphs:
                    texts = paragraph.findall('.//w:t', namespaces)
                    paragraph_text = ''.join([t.text for t in texts if t.text])
                    if paragraph_text:
                        text.append(paragraph_text)
                
                return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error en extracción básica de DOCX: {e}")
            return None
    
    def _generate_content(self, prompt: str, use_stream: bool = False, max_tokens: int = None):
        """
        Método interno para generar contenido con Groq
        
        Args:
            prompt: Prompt completo
            use_stream: Si usar streaming o no
            max_tokens: Máximo de tokens a generar (usa GROQ_MAX_TOKENS_DECLARATION por defecto)
        
        Returns/Yields:
            str o Generator de strings
        """
        # Usar valor por defecto si no se especifica
        if max_tokens is None:
            max_tokens = GROQ_MAX_TOKENS_DECLARATION
        try:
            logger.debug(f"Generando contenido con Groq (modelo: {self.model_name})...")
            start_time = time.time()
            
            if use_stream:
                # Modo streaming
                stream = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=max_tokens,
                    stream=True,
                    timeout=self.request_timeout
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content:
                        yield chunk.choices[0].delta.content
                
                elapsed_time = time.time() - start_time
                logger.debug(f"Generación con streaming completada en {elapsed_time:.2f} segundos")
            else:
                # Modo normal
                completion = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=max_tokens,
                    timeout=self.request_timeout
                )
                
                elapsed_time = time.time() - start_time
                logger.debug(f"Generación completada en {elapsed_time:.2f} segundos")
                
                return completion.choices[0].message.content
        
        except Exception as e:
            error_msg = str(e)
            if "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
                logger.error(f"La generación excedió el tiempo límite de {self.request_timeout}s")
            else:
                logger.error(f"Error al generar contenido con Groq: {e}")
            raise Exception(f"Error al generar contenido: {error_msg}")
    
    def _build_prompt(self, questionnaire_text: str) -> str:
        """Construye el prompt para Declaration Letter"""
        prompt = f"""
{self.system_prompt}

{self.declaration_guide}

---

CUESTIONARIO DEL AFECTADO:
{questionnaire_text}

---
"""
        return prompt
    
    def _build_cover_letter_prompt(self, declaration_letter_content: str) -> str:
        """Construye el prompt para Cover Letter"""
        prompt = f"""
{self.cover_letter_system_prompt}

{self.cover_letter_structure}

---

DECLARATION LETTER DEL SOBREVIVIENTE:
{declaration_letter_content}

---
"""
        return prompt
    
    def generate_declaration_letter(self, questionnaire_text: str) -> Optional[str]:
        """Genera una declaration letter basada en el cuestionario"""
        full_prompt = self._build_prompt(questionnaire_text)
        return self._generate_content(full_prompt, use_stream=False, max_tokens=GROQ_MAX_TOKENS_DECLARATION)
    
    def generate_declaration_letter_stream(self, questionnaire_text: str) -> Generator[str, None, None]:
        """Genera una declaration letter con streaming"""
        full_prompt = self._build_prompt(questionnaire_text)
        yield from self._generate_content(full_prompt, use_stream=True, max_tokens=GROQ_MAX_TOKENS_DECLARATION)
    
    def generate_cover_letter(self, declaration_letter_content: str) -> Optional[str]:
        """Genera un Cover Letter basado en el Declaration Letter"""
        if not self.cover_letter_system_prompt or not self.cover_letter_structure:
            logger.error("Archivos XML de Cover Letter no cargados")
            return None
        
        full_prompt = self._build_cover_letter_prompt(declaration_letter_content)
        return self._generate_content(full_prompt, use_stream=False, max_tokens=GROQ_MAX_TOKENS_COVER)
    
    def generate_cover_letter_stream(self, declaration_letter_content: str) -> Generator[str, None, None]:
        """Genera un Cover Letter con streaming"""
        if not self.cover_letter_system_prompt or not self.cover_letter_structure:
            logger.error("Archivos XML de Cover Letter no cargados")
            raise Exception("Cover Letter XML files not loaded")
        
        full_prompt = self._build_cover_letter_prompt(declaration_letter_content)
        yield from self._generate_content(full_prompt, use_stream=True, max_tokens=GROQ_MAX_TOKENS_COVER)


# FUNCIONES DE UTILIDAD
def create_ai_processor(api_key: str, model_name: str, request_timeout: int = 300) -> Optional[GeminiAIProcessor]:
    """
    Crea y configura un procesador de IA (Gemini por defecto, para compatibilidad)
    
    Args:
        api_key: API key de Google Gemini
        model_name: Nombre del modelo
        request_timeout: Timeout en segundos para las solicitudes (default: 300s = 5 minutos)
    
    Returns:
        GeminiAIProcessor o None si hay error
    """
    try:
        processor = GeminiAIProcessor(api_key, model_name, request_timeout)
        
        # Obtener rutas de archivos XML
        base_path = Path(__file__).parent.parent
        
        # Archivos XML de Declaration Letter
        system_prompt_path = base_path / "DeclarationLetter" / "SystemPrompt.xml"
        declaration_path = base_path / "DeclarationLetter" / "Declaration.xml"
        
        # Archivos XML de Cover Letter
        cover_letter_system_prompt_path = base_path / "CoverLetter" / "SystemPrompt.xml"
        cover_letter_structure_path = base_path / "CoverLetter" / "CoverLetterStructure.xml"
        
        # Cargar archivos XML de Declaration Letter
        if not processor.load_xml_files(str(system_prompt_path), str(declaration_path)):
            return None
        
        # Cargar archivos XML de Cover Letter
        if not processor.load_cover_letter_xml_files(str(cover_letter_system_prompt_path), str(cover_letter_structure_path)):
            logger.warning("No se pudieron cargar archivos XML de Cover Letter")
            # No falla la creación del procesador, solo advierte
        
        return processor
    
    except Exception as e:
        logger.error(f"Error al crear procesador de IA: {e}")
        return None
